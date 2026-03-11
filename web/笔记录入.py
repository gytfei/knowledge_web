# -*- coding: utf-8 -*-
from __future__ import annotations

import os
import re
import sqlite3
import traceback
from dataclasses import dataclass
from pathlib import Path
from datetime import date
# import datetime
from pathlib import Path
import streamlit as st
from difflib import SequenceMatcher
from collections import OrderedDict
from streamlit_paste_button import paste_image_button
from docx import Document
from docx.shared import Inches
from PIL import Image
import mammoth
import platform
import time
# 🔥 追加写入 updated.txt
from datetime import datetime
from streamlit_paste_button import paste_image_button
import io
import re
# st.write("当前操作系统:", platform.system())
# =========================================================
# 0) 相对路径配置（项目根目录 = web/ 的上一级）
# =========================================================
BASE_DIR = Path(__file__).resolve().parents[1]          # your_project/
DATA_DIR = BASE_DIR / "data"
APP_STATE_DIR = DATA_DIR / "app_state"
DATABASE_ROOT = DATA_DIR / "Database"
DATABASE_FILE_DIR = DATABASE_ROOT / "File"

APP_STATE_DIR.mkdir(parents=True, exist_ok=True)
DATABASE_ROOT.mkdir(parents=True, exist_ok=True)
DATABASE_FILE_DIR.mkdir(parents=True, exist_ok=True)

P_UPDATED_TXT = APP_STATE_DIR / "updated.txt"
P_ROOTPATH_TXT = APP_STATE_DIR / "Rootpath.txt"
P_PREPAGE_TXT = APP_STATE_DIR / "Prepage.txt"
P_LAST_TITLE_TXT = APP_STATE_DIR / "last_title.txt"
P_NEWKEYWORD_TXT = APP_STATE_DIR / "newkeyword.txt"
P_LABEL_TXT = APP_STATE_DIR / "label.txt"
P_TEMPLATE_DOCX = APP_STATE_DIR / "template.docx"
P_TEMP_PNG = APP_STATE_DIR / "temp.png"

P_DATABASE_DB = DATABASE_FILE_DIR / "Database.db"

st.markdown("""
<style>

/* 隐藏默认的多页面导航标题 */
section[data-testid="stSidebar"] div[data-testid="stSidebarNav"] > ul {
    margin-top: 10px;
}

/* 改 sidebar 页面文字样式 */
section[data-testid="stSidebar"] div[data-testid="stSidebarNav"] span {
    font-size: 20px !important;
    font-weight: 700 !important;
}

</style>
""", unsafe_allow_html=True)
# =========================================================
# 1) 通用小工具
# =========================================================
def remove_invalid_characters(input_string: str) -> str:
    # 移除无效的XML字符，避免 docx 写入报错
    return re.sub(
        u"[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+",
        "",
        input_string,
    )
def convert_doc_path(doc_path, window_root, ubuntu_root):
    doc_path = Path(str(doc_path).replace("\\", "/"))
    window_root = Path(str(window_root).replace("\\", "/"))

    try:
        relative_part = doc_path.relative_to(window_root)
    except ValueError:
        return str(doc_path)

    return str(Path(ubuntu_root) / relative_part)
def set_txt_state(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(str(value), encoding="utf-8", errors="ignore")

def read_txt_state(path: Path, default: str = "") -> str:
    if not path.exists():
        return default
    return path.read_text(encoding="utf-8", errors="ignore")

def ensure_file(path: Path, default_text: str = "") -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not path.exists():
        path.write_text(default_text, encoding="utf-8")

def sanitize_keyword(s: str) -> str:
    s = re.sub(r'[\/:*?"<>|\n]', "", s)
    s = remove_invalid_characters(s)
    return s.strip()

def similar(a: str, b: str) -> float:
    return SequenceMatcher(None, a, b).ratio()

def extract_string_from_doc_path(original_string: str) -> str:
    # 你的原逻辑：返回 “资料库” 后面的相对部分
    keyword = "资料库"
    idx = original_string.find(keyword)
    if idx == -1:
        return ""
    return original_string[idx + len(keyword):]

def check_dir(p: Path) -> None:
    p.mkdir(parents=True, exist_ok=True)


# =========================================================
# 2) Database.db 结构初始化（PATH / parameter / history）
# =========================================================
def init_database_db(db_path: Path) -> None:
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()

    cur.execute("""
    CREATE TABLE IF NOT EXISTS PATH (
        Database_name TEXT PRIMARY KEY,
        root_path TEXT
    );
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS parameter (
        item TEXT PRIMARY KEY,
        value TEXT
    );
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS history (
        Date TEXT PRIMARY KEY,
        Num TEXT,
        No TEXT
    );
    """)

    # 初始化 Action_score
    cur.execute("SELECT value FROM parameter WHERE item='Action_score'")
    if cur.fetchone() is None:
        cur.execute("INSERT INTO parameter(item, value) VALUES('Action_score', '0')")

    conn.commit()
    conn.close()


def db_fetch_database_names() -> list[str]:
    init_database_db(P_DATABASE_DB)
    conn = sqlite3.connect(P_DATABASE_DB)
    cur = conn.cursor()
    cur.execute("SELECT Database_name FROM PATH ORDER BY Database_name ASC")
    rows = cur.fetchall()
    conn.close()
    names = [r[0] for r in rows]
    return [""] + names
def remove_prefix_before_database(lines):
    new_lines = []
    for line in lines:
        line = line.strip()
        if "资料库/" in line:
            new_part = line.split("资料库/", 1)[1]
            new_lines.append(new_part + "\n")
    return new_lines

def db_get_root_path(database_name: str) -> str:
    if not database_name:
        return ""
    conn = sqlite3.connect(P_DATABASE_DB)
    cur = conn.cursor()
    cur.execute("SELECT root_path FROM PATH WHERE Database_name=?", (database_name,))
    row = cur.fetchone()
    conn.close()
    return row[0] if row else ""

def db_get_root_path_ubuntu(database_name: str) -> str:
    if not database_name:
        return ""
    conn = sqlite3.connect(P_DATABASE_DB)
    cur = conn.cursor()
    cur.execute("SELECT ubuntu_path FROM PATH WHERE Database_name=?", (database_name,))
    row = cur.fetchone()
    conn.close()
    return row[0] if row else ""
def db_insert_path(database_name: str, root_path: str) -> None:
    init_database_db(P_DATABASE_DB)
    conn = sqlite3.connect(P_DATABASE_DB)
    cur = conn.cursor()
    cur.execute("""
        INSERT OR REPLACE INTO PATH(Database_name, root_path)
        VALUES(?, ?)
    """, (database_name, root_path))
    conn.commit()
    conn.close()


def db_get_action_score() -> str:
    init_database_db(P_DATABASE_DB)
    conn = sqlite3.connect(P_DATABASE_DB)
    cur = conn.cursor()
    cur.execute("SELECT value FROM parameter WHERE item='Action_score'")
    row = cur.fetchone()
    conn.close()
    return row[0] if row else "0"


def record_history_and_increment() -> int:
    """
    对应你原来的 record_history()：
    - 每次导入成功：Action_score +1
    - 新的一天：归 1，并把昨天的数据写入 history.Num / history.No
    """
    init_database_db(P_DATABASE_DB)
    today = date.today().isoformat()
    newday = 0

    conn = sqlite3.connect(P_DATABASE_DB)
    cur = conn.cursor()

    # Today_date
    cur.execute("SELECT value FROM parameter WHERE item='Today_date'")
    row = cur.fetchone()
    if row is None:
        cur.execute("INSERT INTO parameter(item, value) VALUES('Today_date', ?)", (today,))
        newday = 1
    else:
        if row[0] != today:
            cur.execute("UPDATE parameter SET value=? WHERE item='Today_date'", (today,))
            newday = 1

    # Action_score
    cur.execute("SELECT value FROM parameter WHERE item='Action_score'")
    row = cur.fetchone()
    old_score = int(row[0]) if row and row[0].isdigit() else 0

    if newday == 0:
        new_score = old_score + 1
        cur.execute("UPDATE parameter SET value=? WHERE item='Action_score'", (str(new_score),))
        conn.commit()
        conn.close()
        return new_score

    # newday==1：新的一天
    new_score = 1
    cur.execute("UPDATE parameter SET value=? WHERE item='Action_score'", (str(new_score),))

    # history: 确保今天条目存在
    cur.execute("SELECT Date FROM history WHERE Date=?", (today,))
    if cur.fetchone() is None:
        cur.execute("INSERT INTO history(Date, Num, No) VALUES(?, ?, ?)", (today, "", ""))

    # 更新昨天条目（最近一天 < today）
    cur.execute("""
        SELECT Date, Num, No
        FROM history
        WHERE Date < ?
        ORDER BY Date DESC
        LIMIT 2
    """, (today,))
    entries = cur.fetchall()

    # entries[0]=昨天，entries[1]=前天（如果有）
    if len(entries) >= 1:
        yesterday = entries[0]
        y_date = yesterday[0]

        # 把 old_score 写入昨天 Num
        cur.execute("UPDATE history SET Num=? WHERE Date=?", (str(old_score), y_date))

        # No 递推：如果有前天，按日期差累加；否则 No=1
        if len(entries) == 2:
            before = entries[1]
            last_date = datetime.strptime(yesterday[0], "%Y-%m-%d")
            before_date = datetime.strptime(before[0], "%Y-%m-%d")
            date_diff = (last_date - before_date).days

            if before[2] not in (None, ""):
                new_no = int(before[2]) + date_diff
            else:
                new_no = 2
        else:
            new_no = 1

        cur.execute("UPDATE history SET No=? WHERE Date=?", (str(new_no), y_date))

    conn.commit()
    conn.close()
    return new_score


# =========================================================
# 3) 每个数据库目录结构（Syn.db / 资料库 / File / Lib_path.txt）
# =========================================================
def ensure_db_structure(root_path: Path) -> dict[str, Path]:
    """
    root_path = data/Database/<DB_NAME>
    你原本的结构：<root>/file/Syn.db 和 <root>/资料库
    这里统一用：<root>/File/Syn.db （大小写统一）
    """
    file_dir = root_path / "file"
    lib_dir = root_path / "资料库"
    syn_db = file_dir / "Syn.db"
    lib_path_txt = file_dir / "Lib_path.txt"

    file_dir.mkdir(parents=True, exist_ok=True)
    lib_dir.mkdir(parents=True, exist_ok=True)

    if not syn_db.exists():
        conn = sqlite3.connect(syn_db)
        cur = conn.cursor()
        cur.execute("CREATE TABLE IF NOT EXISTS Syn (content TEXT, syn TEXT)")
        conn.commit()
        conn.close()

    ensure_file(lib_path_txt, "")

    return {
        "root": root_path,
        "file_dir": file_dir,
        "lib_dir": lib_dir,
        "syn_db": syn_db,
        "lib_path_txt": lib_path_txt,
    }


# def rebuild_lib_path_index(lib_dir: Path, lib_path_txt: Path) -> int:
#     """
#     扫描 资料库 下所有 doc/docx，写入 Lib_path.txt
#     同时清理不存在路径
#     """
#     paths: list[str] = []
#     for root, _, files in os.walk(lib_dir):
#         for f in files:
#             if (f.endswith(".doc") or f.endswith(".docx")) and ("~$" not in f):
#                 paths.append(str(Path(root) / f))
#     paths=remove_prefix_before_database(paths)
#     lib_path_txt.write_text("\n".join(paths) + ("\n" if paths else ""), encoding="utf-8", errors="ignore")
#     return len(paths)
def rebuild_lib_path_index(lib_dir: Path, lib_path_txt: Path) -> int:
    """
    扫描资料库 lib_dir 下所有 doc/docx 文件，
    写入 lib_path_txt，并返回文件数量
    """

    paths = set()  # 用 set 自动去重

    for root, _, files in os.walk(lib_dir):
        for f in files:

            # 过滤 Word 临时文件
            if "~$" in f:
                continue

            # 只要 doc / docx
            if not f.lower().endswith((".doc", ".docx")):
                continue

            p = Path(root) / f

            # 转为相对路径（更稳定）
            try:
                rel = p.relative_to(lib_dir)
            except ValueError:
                rel = p

            # 统一路径分隔符
            paths.add(rel.as_posix())

    # 排序，保证索引稳定
    paths_sorted = sorted(paths)

    # 写入文件
    lib_path_txt.write_text("\n".join(paths_sorted), encoding="utf-8")

    return len(paths_sorted)

# def load_lib_paths(lib_path_txt: Path) -> list[str]:
#     if not lib_path_txt.exists():
#         return []
#     lines = lib_path_txt.read_text(encoding="utf-8", errors="ignore").splitlines()
#     valid = [ln.strip() for ln in lines if ln.strip() and Path(ln.strip()).exists()]
#     # # 回写清理
#     lib_path_txt.write_text("\n".join(valid) + ("\n" if valid else ""), encoding="utf-8", errors="ignore")
#     return valid
def load_lib_paths(lib_path_txt: Path, selected_db) -> list[str]:
    if not lib_path_txt.exists():
        return []

    raw_lines = lib_path_txt.read_text(
        encoding="utf-8",
        errors="ignore"
    ).splitlines()

    rp = db_get_root_path(selected_db)

    clean_paths = []

    for line in raw_lines:
        line = line.strip()

        if not line:
            continue

        # 统一路径分隔符
        line = line.replace("\\", "/")

        # 去除可能已有的 资料库 前缀
        if line.startswith("资料库/"):
            line = line[len("资料库/"):]

        full_path = os.path.join(rp, "资料库", line)
        clean_paths.append(full_path)

    return clean_paths

def find_doc_path_by_keyword(lib_paths: list[str], keyword: str) -> str:
    for p in lib_paths:
        if platform.system() == "Linux":
            p = p.replace("\\", "/")
        # st.write("Path(p)=", Path(p))
        # st.write("Path(p).stem.=", Path(p).stem)
        name = Path(p).stem.split(".")[0]
        # st.write("name=", name)
        if name == keyword:
            return p

    return ""


# =========================================================
# 4) Syn.db：相似词检索 / 插入 / 更新
# =========================================================
def syn_get_similar_contents(syn_db: Path, user_string: str) -> list[str]:
    conn = sqlite3.connect(syn_db)
    cur = conn.cursor()
    cur.execute("SELECT syn, content FROM Syn")
    rows = cur.fetchall()
    conn.close()

    scored: list[tuple[str, float]] = []
    for syn, content in rows:
        syn = syn or ""
        content = content or ""
        score = similar(syn, user_string)
        if score > 0.7 or (user_string in syn) or (syn in user_string):
            scored.append((content, score))

    scored.sort(key=lambda x: x[1], reverse=True)
    # 去重并保持顺序
    out = [c for c, _ in scored]
    out = list(OrderedDict.fromkeys(out))
    return out


def syn_insert(syn_db: Path, content: str, syn: str) -> None:
    conn = sqlite3.connect(syn_db)
    cur = conn.cursor()
    cur.execute("INSERT INTO Syn(content, syn) VALUES(?, ?)", (content, syn))
    conn.commit()
    conn.close()


def syn_insert_or_update(syn_db: Path, content: str, syn: str) -> None:
    conn = sqlite3.connect(syn_db)
    cur = conn.cursor()
    cur.execute("SELECT 1 FROM Syn WHERE content=? AND syn=?", (content, syn))
    if cur.fetchone() is None:
        cur.execute("INSERT INTO Syn(content, syn) VALUES(?, ?)", (content, syn))
    conn.commit()
    conn.close()


def syn_rename_content(syn_db: Path, old_content: str, new_content: str) -> None:
    conn = sqlite3.connect(syn_db)
    cur = conn.cursor()
    cur.execute("UPDATE Syn SET content=? WHERE content=?", (new_content, old_content))
    conn.commit()
    conn.close()


# =========================================================
# 5) docx 写入（正文 / 图片）
# =========================================================
def append_text_to_docx(docx_path: Path, text: str) -> None:
    doc = Document(str(docx_path)) if docx_path.exists() else Document()
    p = doc.add_paragraph()
    p.add_run(text)
    doc.save(str(docx_path))


def insert_image_into_docx(image_path: Path, docx_path: Path, size: int, ref_num: str, declare: str, label: str) -> None:
    doc = Document(str(docx_path)) if docx_path.exists() else Document()

    if declare.strip():
        paragraph = doc.add_paragraph()
        s = remove_invalid_characters(declare.strip() + ":")
        paragraph.add_run(s)

    with Image.open(image_path) as img:
        w, h = img.size

    doc.add_picture(str(image_path), width=Inches(w *size / 96), height=Inches(h * size / 96))

    paragraph = doc.add_paragraph()
    if ref_num.strip():
        paragraph.add_run(f"[{ref_num}][{label}]")
    else:
        paragraph.add_run(f"[{label}]")

    doc.save(str(docx_path))


# =========================================================
# 6) Streamlit UI（尽量保持你图里的布局）
# =========================================================
def init_app_state_files():
    ensure_file(P_ROOTPATH_TXT, "")
    ensure_file(P_PREPAGE_TXT, "New_keyword")
    ensure_file(P_LAST_TITLE_TXT, "")
    ensure_file(P_NEWKEYWORD_TXT, "")
    ensure_file(P_LABEL_TXT, "User")
    # template.docx 需要你自己放：data/app_state/template.docx
    # temp.png 不强制存在
    init_database_db(P_DATABASE_DB)

def ui_header():
    username = read_txt_state(P_LABEL_TXT, "User")
    actionscore = db_get_action_score()

    col1, col2, col3 = st.columns([2.2, 1.2, 1.2])
    with col1:
        st.markdown("## 📚 Online Note")
    with col2:
        st.metric("今日计数", actionscore)
    with col3:
        st.metric("当前用户", username)

    st.divider()


def ui_left_panel():
    """
    左侧：数据库选择 + 关键词检索 + 结果列表 + 文档定位信息
    """
    st.markdown("### 数据库 / 关键词")
    db_names = db_fetch_database_names()
    # ===== 同一行布局 =====
    col1, col2, col3 = st.columns([1.0, 1.7, 0.5])

    with col1:
        st.markdown("**选择数据库**")
        selected_db = st.selectbox(
            "选择数据库",
            db_names,
            key="db_select",
            label_visibility="collapsed"
        )

    with col2:
        st.markdown("**关键词**")
        keyword = st.text_input(
            "关键词",
            key="keyword_input",
            label_visibility="collapsed"
        )

    with col3:
        with col3:
            st.markdown(
                """
                <div style="height: 33px;"></div>
                """,
                unsafe_allow_html=True
            )
            search_clicked = st.button(
                "检索",
                key="btn_search",
                use_container_width=True,
                type="primary"
            )

    # selected_db = st.selectbox("选择数据库", db_names, index=0, key="db_select")
    # rp = db_get_root_path(selected_db) if selected_db else ""
    if platform.system() == "Windows":
        rp = db_get_root_path(selected_db)
        window_root = rp
    elif platform.system() == "Linux":
        rp = db_get_root_path(selected_db)
        window_root = rp
        rp = db_get_root_path_ubuntu(selected_db)
        ubuntu_root = rp

    #
    # # 显示 root_path
    # st.text_input("数据库根目录（只读）", value=root_path, disabled=True, key="root_path_display")

    # # 快捷：重建索引
    # if selected_db and root_path:
    #     if st.button("重建资料库索引（Lib_path）", key="btn_rebuild_index"):
    #         paths = ensure_db_structure(Path(root_path))
    #         n = rebuild_lib_path_index(paths["lib_dir"], paths["lib_path_txt"])
    #         st.success(f"索引已重建：共 {n} 个 doc/docx")

    # st.divider()
    # keyword = st.text_input("关键词", value=st.session_state.get("keyword_input", ""), key="keyword_input")

    # 查同义词 / 匹配 content
    if search_clicked:
        # keyword2 = sanitize_keyword(keyword)
        # st.session_state["keyword_input"] = keyword2
        keyword2 = sanitize_keyword(keyword)
        # st.write("原始 keyword:", keyword)
        # st.write("清洗后 keyword2:", keyword2)
        set_txt_state(P_LAST_TITLE_TXT, keyword2)

        if not selected_db:
            st.error("请先选择数据库")
        else:
            if platform.system() == "Windows":
                rp = db_get_root_path(selected_db)
                window_root=rp
            elif platform.system() == "Linux":
                rp = db_get_root_path(selected_db)
                window_root = rp
                rp = db_get_root_path_ubuntu(selected_db)
                ubuntu_root=rp
            else:
                raise RuntimeError(f"Unsupported OS: {platform.system()}")
            # st.write("rp=", rp)
            if not rp:
                st.error("数据库 root_path 为空")
            else:
                paths = ensure_db_structure(Path(rp))
                # st.write("paths[syn_db]=", paths["syn_db"])
                # st.write("keyword2=",keyword2)
                results = syn_get_similar_contents(paths["syn_db"], keyword2)
                # st.write("results478=", results)
                st.session_state["search_results"] = results
                st.session_state["selected_content"] = results[0] if results else ""

                # 如果没有结果：模拟你原来跳转“新建关键词/同义词”
                if not results:
                    set_txt_state(P_NEWKEYWORD_TXT, keyword2)
                    set_txt_state(P_ROOTPATH_TXT, rp)
                    set_txt_state(P_PREPAGE_TXT, "New_keyword")

    results = st.session_state.get("search_results", [])
    selected_content = ""
    if results:
        selected_content = st.selectbox("匹配结果", results, key="content_select")
        st.session_state["selected_content"] = selected_content
    else:
        st.info("暂无结果")

    # 显示当前选中的 docx path（来自 Lib_path.txt）
    doc_rel = ""
    doc_path = ""
    if selected_db and rp and st.session_state.get("selected_content"):
        paths = ensure_db_structure(Path(rp))
        # st.write("paths[lib_path_txt]=", paths["lib_path_txt"])
        lib_paths = load_lib_paths(paths["lib_path_txt"],selected_db)

        # st.write("lib_paths[:10]=", lib_paths[:10])
        # st.write("keword=",  st.session_state["selected_content"])
        doc_path = find_doc_path_by_keyword(lib_paths, st.session_state["selected_content"])
        # st.write("doc_path=", doc_path)
        # st.write("platform.system()",platform.system())
        #
        # st.write("window_root",window_root)
        if platform.system() == "Linux":
            # st.write("ubuntu_root", ubuntu_root)
            doc_path = convert_doc_path(
                doc_path,
                window_root,
                ubuntu_root
            )

            # st.write("doc_path552=", doc_path)

        doc_rel = extract_string_from_doc_path(doc_path) if doc_path else ""
    #     st.write("doc_rel555=", doc_rel)
    # #
    # st.text_input("Word 相对路径（资料库后）", value=doc_rel, disabled=True, key="doc_rel_display")

    return selected_db, rp, doc_path

def ui_left_panel_below(selected_db: str, root_path: str, doc_path: str):


    # ========== 右侧下方：同义词管理 / 新建 Word / 新建数据库 ==========
    tab1, tab2, tab3 = st.tabs(["同义词管理", "新建 Word 文档", "新建数据库"])

    with tab1:
        st.markdown("#### 同义词管理（Syn.db）")

        if not selected_db or not root_path:
            st.info("请先选择数据库")
        else:
            paths = ensure_db_structure(Path(root_path))

            # 🔥 直接绑定左侧 selectbox 的值
            current_content = st.session_state.get("content_select", "")

            st.text_input(
                "当前 Content",
                disabled=True,
                key="syn_current_content"
            )

            new_syn = st.text_input(
                "添加一个 Syn（同义词）",
                value="",
                key="syn_new_syn"
            )

            if st.button("添加 Syn", key="btn_syn_add"):
                if not current_content:
                    st.error("当前 Content 为空，请先检索并选择一个 content")

                elif not new_syn.strip():
                    st.error("Syn 不能为空")

                else:
                    syn_insert_or_update(
                        paths["syn_db"],
                        current_content,
                        new_syn.strip()
                    )

                    # 🔥 写入 updated.txt（只写数据库名称）
                    P_UPDATED_TXT.parent.mkdir(parents=True, exist_ok=True)
                    with open(P_UPDATED_TXT, "a", encoding="utf-8") as f:
                        f.write(f"{selected_db}\n")

                    st.success("已添加")

            st.markdown("----")

            rename_to = st.text_input(
                "将 Content 重命名为",
                value="",
                key="syn_rename_to"
            )

            if st.button("执行重命名 Content", key="btn_syn_rename"):
                if not current_content:
                    st.error("当前 Content 为空")
                elif not rename_to.strip():
                    st.error("新名字不能为空")
                else:
                    syn_rename_content(
                        paths["syn_db"],
                        current_content,
                        rename_to.strip()
                    )

                    syn_insert_or_update(
                        paths["syn_db"],
                        rename_to.strip(),
                        rename_to.strip()
                    )

                    st.success("重命名完成（Syn.db 已更新）")

                    # 清空搜索结果，强制刷新
                    st.session_state["search_results"] = []
                    st.session_state["content_select"] = ""
                    st.rerun()

    with tab2:
        st.markdown("#### 新建 Word（从模板复制）")

        if not selected_db or not root_path:
            st.info("请先选择数据库")
        else:
            paths = ensure_db_structure(Path(root_path))
            lib_dir = paths["lib_dir"]

            st.caption(f"资料库目录：{lib_dir}")

            # ===============================
            # 第一行：Word 名称 + 新文件夹名称
            # ===============================
            col_name1, col_name2 = st.columns(2)

            with col_name1:
                new_name = st.text_input(
                    "新 Word 名称（不含 .docx）",
                    value="",
                    key="new_doc_name"
                )

            with col_name2:
                new_folder_name = st.text_input(
                    "新文件夹名称（可空）",
                    value="",
                    key="new_folder_name"
                )

            # ===============================
            # 扫描现有子目录
            # ===============================
            def get_subfolders(base_dir: Path) -> list[str]:
                base_dir = base_dir.resolve()
                folders = [""]  # 空表示根目录

                for root, dirs, _ in os.walk(base_dir):
                    root_path = Path(root)
                    for d in dirs:
                        full_path = (root_path / d).resolve()
                        rel_path = full_path.relative_to(base_dir)
                        folders.append(str(rel_path))

                return sorted(set(folders))

            subfolders = get_subfolders(lib_dir)

            subdir_select = st.selectbox(
                "选择放入资料库的子文件夹",
                subfolders,
                key="new_doc_subdir"
            )

            # ===============================
            # 创建按钮
            # ===============================
            if st.button("创建 Word", key="btn_create_doc"):
                if not P_TEMPLATE_DOCX.exists():
                    st.error("未找到模板：data/app_state/template.docx（请放入模板文件）")

                elif not new_name.strip():
                    st.error("Word 名称不能为空")

                else:
                    # 🔥 构造最终目录
                    final_dir = lib_dir

                    if subdir_select:
                        final_dir = final_dir / subdir_select

                    if new_folder_name.strip():
                        final_dir = final_dir / new_folder_name.strip()

                    # 创建目录（如果不存在）
                    final_dir.mkdir(parents=True, exist_ok=True)

                    dst = final_dir / f"{new_name.strip()}.docx"

                    if dst.exists():
                        st.error("已存在同名 docx")
                    else:
                        dst.write_bytes(P_TEMPLATE_DOCX.read_bytes())
                        st.success(f"创建成功：{dst}")
                        # 🔥 写入 updated.txt
                        P_UPDATED_TXT.parent.mkdir(parents=True, exist_ok=True)

                        with open(P_UPDATED_TXT, "a", encoding="utf-8") as f:
                            f.write(f"{selected_db}\n")
                            f.write(str(dst)+"\n")


                        # 重建索引
                        n = rebuild_lib_path_index(lib_dir, paths["lib_path_txt"])
                        st.toast(f"索引更新：{n} files", icon="📌")

                        # 同义词录入
                        syn_insert_or_update(
                            paths["syn_db"],
                            new_name.strip(),
                            new_name.strip()
                        )

                        st.toast("已写入同义词：content=syn=new_name", icon="🧠")

    with tab3:
        st.markdown("#### 新建数据库（相对路径自动创建在 data/Database/ 下）")
        db_new_name = st.text_input("新数据库名称", value="", key="db_new_name")
        if st.button("创建数据库", key="btn_create_db"):
            if not db_new_name.strip():
                st.error("数据库名不能为空")
            else:
                root = DATABASE_ROOT / db_new_name.strip()
                paths = ensure_db_structure(root)
                # 写入 PATH 表
                db_insert_path(db_new_name.strip(), str(root))
                # 建议初次建立索引
                rebuild_lib_path_index(paths["lib_dir"], paths["lib_path_txt"])
                st.success(f"数据库创建完成：{root}")

def docx_to_html(docx_path: Path) -> str:
    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value  # HTML 字符串
    return html

def get_subfolders(base_dir: Path) -> list[str]:
    folders = [""]
    for root, dirs, _ in os.walk(base_dir):
        for d in dirs:
            full = Path(root) / d
            rel = full.relative_to(base_dir)
            folders.append(str(rel))
    return sorted(folders)

def ui_right_panel(selected_db: str, root_path: str, doc_path: str):
    st.divider()

    if "content_select" in st.session_state:
        st.session_state["syn_current_content"] = st.session_state["content_select"]

    st.markdown("### 内容录入")

    colA, colB = st.columns([1, 3])

    with colA:
        ref_num = st.text_input(
            "引用源",
            value=st.session_state.get("ref_num", ""),
            key="ref_num"
        )

    with colB:
        declare = st.text_input(
            "小标题",
            value=st.session_state.get("declare", ""),
            key="declare"
        )

    col_opt1, col_opt2 = st.columns([2, 1], gap="small")

    with col_opt1:
        hold = st.checkbox(
            "保持小标题",
            value=st.session_state.get("hold", False),
            key="hold"
        )

    # ========= 两个按钮对齐 =========
    with col_opt2:
        btn_col1, btn_col2 = st.columns(2, gap="small")

        # 🔥 压缩连续回车为一个
        with btn_col1:
            if st.button("去掉多余回车", key="btn_remove_double_enter", use_container_width=True):
                text = st.session_state.get("editor_text", "")
                if text:
                    # 统一换行符为 \n
                    text = text.replace("\r\n", "\n").replace("\r", "\n")

                    # 将连续两个以上换行压缩为一个
                    text = re.sub(r"\n{2,}", "\n", text)

                    st.session_state["editor_text"] = text

        # 🔥 删除所有回车
        with btn_col2:
            if st.button("去掉回车符", key="btn_remove_enter", use_container_width=True):
                text = st.session_state.get("editor_text", "")
                if text:
                    text = text.replace("\r\n", "").replace("\n", "").replace("\r", "")
                    st.session_state["editor_text"] = text

    # ========= 正文编辑 =========
    st.markdown(
        "<div style='font-size:22px;font-weight:700;margin-bottom:-15px;'>正文编辑区</div>",
        unsafe_allow_html=True
    )

    if "editor_text" not in st.session_state:
        st.session_state.editor_text = ""

    editor_text = st.text_area(
        " ",
        key="editor_text",
        height=180
    )

    # ========= 底部按钮 =========
    col_btn1, col_btn2, col_btn3 = st.columns([1.5, 2.5, 2], gap="small")

    with col_btn1:
        if st.button("清空", key="btn_clear_editor", use_container_width=True):
            st.session_state["editor_text"] = ""

    with col_btn2:
        label = read_txt_state(P_LABEL_TXT, "User")

        if st.button(
                "保存笔记",
                key="btn_import_docx",
                use_container_width=True,
                type="primary"
        ):
            try:
                if not selected_db or not root_path:
                    st.error("请先选择数据库")
                    return

                if not doc_path:
                    st.error("未定位到 Word 文档路径")
                    return

                if not Path(doc_path).exists():
                    st.error("docx 路径不存在")
                    return

                content = (st.session_state.get("editor_text") or "").strip()
                if not content:
                    st.error("正文编辑区为空")
                    return

                ref = (ref_num or "").strip()
                dec = (declare or "").strip()

                if dec:
                    s = f"{dec}:{content}"
                else:
                    s = content

                s = remove_invalid_characters(s)

                if "\n" in s:
                    s = "{" + s + "}"

                s = s + f"[{ref}]" + f"[{label}]"

                append_text_to_docx(Path(doc_path), s)

                P_UPDATED_TXT.parent.mkdir(parents=True, exist_ok=True)

                with open(P_UPDATED_TXT, "a", encoding="utf-8") as f:
                    f.write(f"{selected_db}\n")
                    f.write(f"{doc_path}\n")

                new_score = record_history_and_increment()
                st.toast(f"Action Score = {new_score}", icon="✅")

                del st.session_state["editor_text"]

                st.rerun()

            except Exception:
                st.error("写入失败：\n" + traceback.format_exc())

    with col_btn3:
        if st.button("网页打开", key="btn_open_html", use_container_width=True):
            if not doc_path or not Path(doc_path).exists():
                st.error("未找到 Word 文件")
            else:
                st.session_state["preview_doc_path"] = doc_path
                st.switch_page("pages/文件查看.py")

def ui_right_panel_below(selected_db: str, root_path: str, doc_path: str):
    st.divider()
    st.markdown("### 图片保存")

    # 初始化图片状态
    if "image_to_use" not in st.session_state:
        st.session_state.image_to_use = None

    # ====== 两行两列布局 ======
    row1_col1, row1_col2 = st.columns(2)
    row2_col1, row2_col2 = st.columns(2)

    # =============================
    # 左上：上传 PNG
    # =============================
    with row1_col1:
        img_file = st.file_uploader(
            "上传 PNG/JPG",
            type=["png", "jpg", "jpeg"],
            key="img_uploader"
        )

        if img_file is not None:
            st.session_state.image_to_use = Image.open(img_file).convert("RGB")

    # =============================
    # 右上：图片缩放
    # =============================
    with row1_col2:
        size = st.selectbox(
            "图片缩放（size=1, 1/2, 1/3）",
            [0.3, 0.6, 1.0],
            index=0,
            key="img_size"
        )

    # =============================
    # 右下：插入按钮
    # =============================
    with row2_col2:

        if st.button("插入图片", key="btn_insert_img", use_container_width=True):
            try:
                image_to_use = st.session_state.image_to_use

                if image_to_use is None:
                    st.error("请先上传或粘贴图片")
                    return

                if not selected_db or not root_path or not doc_path:
                    st.error("请先选择数据库并定位到 Word")
                    return

                P_TEMP_PNG.parent.mkdir(parents=True, exist_ok=True)
                image_to_use.save(P_TEMP_PNG)

                insert_image_into_docx(
                    image_path=P_TEMP_PNG,
                    docx_path=Path(doc_path),
                    size=float(size),
                    ref_num=st.session_state.get("ref_num", ""),
                    declare=st.session_state.get("declare", ""),
                    label=read_txt_state(P_LABEL_TXT, "User"),
                )

                st.success("图片插入成功")

                new_score = record_history_and_increment()
                st.toast(f"Action Score = {new_score}", icon="🖼️")

                # if not st.session_state.get("hold", False):
                #     st.session_state["declare"] = ""
                if not st.session_state.get("hold", False):
                    del st.session_state["declare"]
                    st.rerun()
            except Exception as e:
                st.error(f"插入失败: {e}")
    # =============================
    # 左下：粘贴图片
    # =============================
    with row2_col1:

        pasted = paste_image_button(
            label="点击这里或 Ctrl+V 粘贴图片",
            key="paste_image"
        )

        if (
                pasted is not None
                and hasattr(pasted, "image_data")
                and pasted.image_data is not None
        ):
            try:
                st.session_state.image_to_use = pasted.image_data.convert("RGB")
            except Exception as e:
                st.error(f"粘贴图片解析失败: {e}")
    # =============================
    # 统一图片预览（底部）
    # =============================
    if st.session_state.image_to_use is not None:
        st.image(st.session_state.image_to_use, caption="当前图片")




def main():
    # st.set_page_config(page_title="共享笔记本", layout="centered")
    st.set_page_config(page_title="共享笔记本", layout="wide")
    
    st.markdown("""
    <style>

    /* ===== 1️⃣ 页面整体顶部距离 ===== */
    .block-container {
        padding-top: 2.5rem !important;   /* 再往上收 */
        padding-bottom: 0rem !important;
    }

    /* ===== 2️⃣ 去掉标题默认 margin ===== */
    h1, h2, h3 {
        margin-top: 0rem !important;
        margin-bottom: 0.0rem !important;
    }

    /* ===== 3️⃣ 压缩 columns 内部高度 ===== */
    div[data-testid="column"] {
        padding-top: 0rem !important;
        padding-bottom: 0rem !important;
    }

    /* ===== 4️⃣ 压缩 metric 组件 ===== */
    [data-testid="stMetric"] {
        padding-top: 0rem !important;
        padding-bottom: 0rem !important;
        margin-top: 0rem !important;
        margin-bottom: 0rem !important;
    }

    /* ===== 5️⃣ 横线贴近上方内容 ===== */
    hr {
        margin-top: 0.0rem !important;
        margin-bottom: 0.0rem !important;
    }
    
    /* 🔥 彻底压缩 text_input 上下间距 */
    div[data-testid="stTextInput"] {
        margin-top: -18px !important;
        margin-bottom: -4px !important;
    }
    
    /* 🔥 压缩 label 到输入框的距离 */
    div[data-testid="stTextInput"] label {
        margin-bottom: 0px !important;
        padding-bottom: 0px !important;
    }
    
    /* 🔥 压缩 columns 内部块间距 */
    div[data-testid="column"] > div {
        margin-bottom: -12px !important;
    }
    
    /* 🔥 再压缩整体 block 间距 */
    .block-container div {
        line-height: 1.1 !important;
    }
    div[data-testid="stButton"] button {
        font-weight: 700 !important;
    }

    </style>
    """, unsafe_allow_html=True)


    st.markdown("""
    <style>
    /* =============================
       让 Selectbox / TextInput / Button 真正同高同基线
       ============================= */

    /* 1) 统一外层控件块的下方空隙（可选，防止上下跳） */
    div[data-testid="stSelectbox"], 
    div[data-testid="stTextInput"], 
    div[data-testid="stButton"]{
        margin-top: 0rem !important;
        margin-bottom: 0rem !important;
    }

    /* 2) Selectbox：控制框（BaseWeb Select control） */
    div[data-testid="stSelectbox"] div[data-baseweb="select"] > div {
        min-height: 42px !important;
        height: 42px !important;
        padding-top: 0px !important;
        padding-bottom: 0px !important;
        display: flex !important;
        align-items: center !important;   /* 垂直居中 */
    }

    /* 3) Selectbox：内部文字行高（避免视觉偏移） */
    div[data-testid="stSelectbox"] div[data-baseweb="select"] span {
        line-height: 42px !important;
    }

    /* 4) TextInput：输入框本体 */
    div[data-testid="stTextInput"] input {
        min-height: 42px !important;
        height: 42px !important;
        padding-top: 0px !important;
        padding-bottom: 0px !important;
        line-height: 42px !important;
        box-sizing: border-box !important;
    }

    /* 5) Button：统一高度（覆盖 secondary / primary 等） */
    div[data-testid="stButton"] button {
        min-height: 42px !important;
        height: 42px !important;
        padding-top: 0px !important;
        padding-bottom: 0px !important;
        line-height: 42px !important;
    }

    </style>
    """, unsafe_allow_html=True)

    init_app_state_files()

    ui_header()

    left, right = st.columns([1.35, 1.35], gap="large")

    with left:
        selected_db, root_path, doc_path = ui_left_panel()
        ui_right_panel(selected_db, root_path, doc_path)
    with right:
        ui_left_panel_below(selected_db, root_path, doc_path)
        ui_right_panel_below(selected_db, root_path, doc_path)

    # with left:

if __name__ == "__main__":
    main()
